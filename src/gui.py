import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import threading
from pathlib import Path
import os
from dotenv import load_dotenv
# from docx2pdf import convert  # Optional for PDF conversion
# from pdf2image import convert_from_path  # Optional, wird nicht benötigt
from PIL import Image, ImageTk
import time
import logging
from datetime import datetime
import pythoncom
import glob
import docx
import json
import re

from .ai_service import AIServiceManager
from .word_generator import WordDocumentGenerator


class GuiLogHandler(logging.Handler):
    def __init__(self, gui_log_func):
        super().__init__()
        self.gui_log_func = gui_log_func

    def emit(self, record):
        msg = self.format(record)
        self.gui_log_func(msg)

class ZetaProposerGUI:
    CONFIG_PATH = "config.json"
    def __init__(self, root):
        self.root = root
        self.root.title("Zeta Proposer - Technical Concept Generator")
        self.root.geometry("800x700")
        
        # Load environment variables (for backward compatibility)
        load_dotenv()
        
        # Initialize default values first
        self.log_widget = None
        self.logger = None
        self.logfile_path = None
        self.selected_template = None
        self.ai_provider_var = tk.StringVar(value="openai")
        self.openai_api_key = ""
        self.openai_model = "gpt-4o"
        self.ollama_url = "http://localhost:11434"
        self.ollama_model = "llama3"
        self.alignment_threshold = 0.6  # Defaultwert
        self.output_directory = "output/docx"  # Defaultwert
        self.json_output_directory = "output/json"  # Defaultwert for JSON files
        self.initiator = ""  # Defaultwert
        self.cancel_requested = False  # Für Abbrechen-Button
        
        # Load configuration (this will override defaults)
        self.load_config()
        
        # Initialize services with configured output directory
        self.ai_service = AIServiceManager()
        self.word_generator = WordDocumentGenerator(self.output_directory)
        
        # Set template if available
        if self.selected_template:
            self.word_generator.set_template(self.selected_template)
        
        self.setup_ui()
        # Beim Schließen speichern
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
    def on_closing(self):
        # Speichere aktuelle Einstellungen (inkl. Threshold)
        self.save_config()
        self.root.destroy()
        
    def show_documents_overview(self):
        """Zeigt eine Übersicht aller erstellten docx-Dokumente (rekursiv in Unterordnern) und ermöglicht das Öffnen per Klick."""
        import glob
        import platform
        import subprocess
        from tkinter import Toplevel, Listbox, Button, END, Scrollbar, RIGHT, Y, LEFT, BOTH, Label
        from pathlib import Path
        
        # Verwende den aktuell konfigurierten Output-Ordner
        output_dir = Path(self.output_directory)
        if not output_dir.exists():
            messagebox.showwarning("Warnung", f"Output-Ordner existiert nicht: {output_dir}")
            return
        
        # Rekursive Suche nach allen docx-Dateien in Unterordnern
        files = sorted(output_dir.glob("*/*.docx"), key=os.path.getmtime, reverse=True)
        win = Toplevel(self.root)
        win.title(f"Erstellte Dokumente - {output_dir}")
        win.geometry("600x450")
        
        # Zeige den aktuellen Output-Ordner an
        folder_label = Label(win, text=f"Output-Ordner: {output_dir}", font=("Arial", 9))
        folder_label.pack(pady=(5, 0))
        lb = Listbox(win, width=80)
        lb.pack(side=LEFT, fill=BOTH, expand=True)
        sb = Scrollbar(win)
        sb.pack(side=RIGHT, fill=Y)
        lb.config(yscrollcommand=sb.set)
        sb.config(command=lb.yview)
        for f in files:
            lb.insert(END, str(f.relative_to(output_dir)))
        def open_selected(event=None):
            sel = lb.curselection()
            if not sel:
                return
            file_path = output_dir / lb.get(sel[0])
            file_path = file_path.resolve()
            try:
                if platform.system() == "Windows":
                    os.startfile(str(file_path))
                elif platform.system() == "Darwin":
                    subprocess.Popen(["open", str(file_path)])
                else:
                    subprocess.Popen(["xdg-open", str(file_path)])
            except Exception as e:
                messagebox.showerror("Fehler", f"Konnte Datei nicht öffnen: {e}")
        lb.bind('<Double-Button-1>', open_selected)
        
    def setup_ui(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=0)  # Label-Spalte
        main_frame.columnconfigure(1, weight=1)  # Eingabefeld-Spalte
        # Title
        title_label = ttk.Label(main_frame, text="Zeta Proposer", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        # Settings button
        settings_btn = ttk.Button(main_frame, text="Einstellungen", command=self.open_settings)
        settings_btn.grid(row=0, column=2, padx=(10, 0), pady=5, sticky="e")
        # Project Name
        ttk.Label(main_frame, text="Project Name:").grid(row=1, column=0, sticky="nw", pady=5)
        self.project_name_var = tk.StringVar()
        self.project_name_entry = ttk.Entry(main_frame, textvariable=self.project_name_var)
        self.project_name_entry.grid(row=1, column=1, sticky="ew", padx=(10, 0), pady=5)
        self.project_name_entry.insert(0, "Enter project name here...")
        self.project_name_entry.bind('<FocusIn>', lambda e: self.on_project_name_focus_in())
        self.project_name_entry.bind('<FocusOut>', lambda e: self.on_project_name_focus_out())
        # Upwork Link
        ttk.Label(main_frame, text="Upwork Link:").grid(row=2, column=0, sticky="nw", pady=5)
        self.upwork_link_var = tk.StringVar()
        self.upwork_link_entry = ttk.Entry(main_frame, textvariable=self.upwork_link_var)
        self.upwork_link_entry.grid(row=2, column=1, sticky="ew", padx=(10, 0), pady=5)
        # Project Description
        ttk.Label(main_frame, text="Project Description:").grid(row=3, column=0, sticky="nw", pady=5)
        self.description_text = scrolledtext.ScrolledText(main_frame, height=10)
        self.description_text.grid(row=3, column=1, sticky="nsew", padx=(10, 0), pady=5)
        main_frame.rowconfigure(3, weight=1)
        # Buttons frame
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=20)
        self.generate_btn = ttk.Button(button_frame, text="Generate Technical Concept", command=self.generate_concept)
        self.generate_btn.pack(side=tk.LEFT, padx=(0, 10))
        self.cancel_btn = ttk.Button(button_frame, text="Abbrechen", command=self.cancel_generation, state="disabled")
        self.cancel_btn.pack(side=tk.LEFT, padx=(0, 10))
        load_btn = ttk.Button(button_frame, text="Load from File", command=self.load_from_file)
        load_btn.pack(side=tk.LEFT, padx=(0, 10))
        generate_json_btn = ttk.Button(button_frame, text="Generate from Specification", command=self.generate_json_from_specification)
        generate_json_btn.pack(side=tk.LEFT, padx=(0, 10))
        self.progress_var = tk.StringVar(value="Ready")
        progress_label = ttk.Label(button_frame, textvariable=self.progress_var)
        progress_label.pack(side=tk.LEFT, padx=(20, 0))
        # Status bar
        self.status_var = tk.StringVar(value="Ready")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN)
        status_bar.grid(row=5, column=0, columnspan=2, sticky="ew", pady=(10, 0))
        # Log window (initial ausgeblendet)
        self.log_visible = False
        self.log_label = ttk.Label(main_frame, text="Program Log:")
        self.log_widget = scrolledtext.ScrolledText(main_frame, height=8, state="disabled")
        # Log-Label und Log-Widget werden NICHT gegridet (also nicht sichtbar) beim Start
        self.toggle_log_btn = ttk.Button(main_frame, text="Log anzeigen", command=self.toggle_log)
        self.toggle_log_btn.grid(row=6, column=0, sticky=tk.W, pady=(0, 10))
        view_docs_btn = ttk.Button(main_frame, text="Erstellte Dokumente anzeigen", command=self.show_documents_overview)
        view_docs_btn.grid(row=6, column=1, sticky=tk.W, pady=(0, 10))
        self.project_name_entry.focus()

    def on_project_name_focus_in(self):
        """Handle focus in event for project name entry"""
        if self.project_name_var.get() == "Enter project name here...":
            self.project_name_var.set("")
            
    def on_project_name_focus_out(self):
        """Handle focus out event for project name entry"""
        if not self.project_name_var.get().strip():
            self.project_name_var.set("Enter project name here...")
        
    def setup_logger(self):
        """Set up a new logger and log file for each generation run"""
        # Immer im output-Ordner des aktuellen Projekts, nicht im konfigurierten Output-Ordner
        logs_dir = Path("output") / "logs"
        logs_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.logfile_path = str(logs_dir / f"zeta_log_{timestamp}.log")
        # Remove all handlers before reconfiguring logging to avoid duplicate logs
        for handler in logging.root.handlers[:]:
            logging.root.removeHandler(handler)
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s %(levelname)s %(message)s',
            handlers=[
                logging.FileHandler(self.logfile_path, encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        # Remove all handlers from root logger before adding GUI handler
        for handler in logging.getLogger().handlers[:]:
            if isinstance(handler, GuiLogHandler):
                logging.getLogger().removeHandler(handler)
        # Add GUI log handler
        gui_handler = GuiLogHandler(self.log_message)
        gui_handler.setLevel(logging.INFO)
        logging.getLogger().addHandler(gui_handler)
        self.log_message(f"[LOGGING STARTED] Logfile: {self.logfile_path}")

    def on_ai_provider_change(self, event=None):
        """Handle AI provider change"""
        provider = self.ai_provider_var.get()
        self.status_var.set(f"AI Provider changed to: {provider}")
        
    def load_config(self):
        try:
            with open(self.CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            self.ai_provider_var.set(cfg.get("ai_provider", "openai"))
            self.selected_template = cfg.get("selected_template")
            self.openai_api_key = cfg.get("openai_api_key", "")
            self.openai_model = cfg.get("openai_model", "gpt-4o")
            self.ollama_url = cfg.get("ollama_url", "http://localhost:11434")
            self.ollama_model = cfg.get("ollama_model", "llama3")
            self.alignment_threshold = float(cfg.get("alignment_threshold", 0.6))
            self.output_directory = cfg.get("output_directory", "output/docx")
            self.json_output_directory = cfg.get("json_output_directory", "output/json")
            self.initiator = cfg.get("initiator", "")
        except FileNotFoundError:
            # Erstelle Standard-Konfiguration wenn Datei nicht existiert
            self._create_default_config()
        except Exception:
            # Bei anderen Fehlern auch Standard-Konfiguration erstellen
            self._create_default_config()

    def _create_default_config(self):
        """Erstelle eine Standard-Konfiguration wenn config.json nicht existiert"""
        default_config = {
            "ai_provider": "openai",
            "selected_template": None,
            "openai_api_key": "",
            "openai_model": "gpt-4o",
            "ollama_url": "http://localhost:11434",
            "ollama_model": "llama3",
            "alignment_threshold": 0.6,
            "output_directory": "output/docx",
            "json_output_directory": "output/json",
            "initiator": ""
        }
        try:
            with open(self.CONFIG_PATH, "w", encoding="utf-8") as f:
                json.dump(default_config, f, indent=2)
            print(f"Created default config file: {self.CONFIG_PATH}")
        except Exception as e:
            print(f"Could not create default config file: {e}")

    def _reinitialize_services(self):
        """Reinitialize services with current output directory"""
        self.word_generator = WordDocumentGenerator(self.output_directory)
        
        # Set template if available
        if hasattr(self, 'selected_template') and self.selected_template:
            self.word_generator.set_template(self.selected_template)
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Template set in word generator: %s", self.selected_template)

    def save_config(self):
        cfg = {
            "ai_provider": self.ai_provider_var.get(),
            "selected_template": self.selected_template,
            "openai_api_key": self.openai_api_key,
            "openai_model": self.openai_model,
            "ollama_url": self.ollama_url,
            "ollama_model": self.ollama_model,
            "alignment_threshold": self.alignment_threshold,
            "output_directory": self.output_directory,
            "json_output_directory": self.json_output_directory,
            "initiator": self.initiator
        }
        with open(self.CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2)

    def open_settings(self):
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Einstellungen")
        settings_window.geometry("600x700")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        # Create main frame with scrollbar
        main_frame = ttk.Frame(settings_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Create canvas for scrolling
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Info box for placeholders
        info_text = (
            "Verfügbare Platzhalter für das Technical Concept (im Template verwenden):\n"
            "  {{system_scope}}  → System scope and boundaries\n"
            "  {{architecture_tech_stack}}  → Architecture and technology stack\n"
            "  {{external_interfaces}}  → System-external interfaces and integrations\n"
            "  {{ci_cd}}  → CI/CD Pipelines\n"
            "  {{testing_concept}}  → Specific testing concept\n"
            "  {{deployment_operation}}  → Deployment and Operation environment\n"
            "  {{ux_ui}}  → UX/UI design and prototyping\n"
            "\nJeder Platzhalter wird durch den jeweiligen Abschnitt ersetzt. Die Formatierung folgt den Überschriften- und Text-Styles des Templates."
        )
        info_label = tk.Label(scrollable_frame, text=info_text, justify=tk.LEFT, anchor="w", bg="#f0f0f0", relief=tk.SUNKEN, wraplength=540)
        info_label.pack(fill=tk.X, pady=(0, 15))
        
        # AI Provider
        ttk.Label(scrollable_frame, text="AI Provider:").pack(anchor=tk.W)
        ai_provider_combo = ttk.Combobox(scrollable_frame, textvariable=self.ai_provider_var, values=["openai", "ollama"], state="readonly")
        ai_provider_combo.pack(fill=tk.X, pady=(0, 10))
        
        # Configuration frame
        config_frame = ttk.Frame(scrollable_frame)
        config_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Variables for settings
        openai_api_key_var = tk.StringVar(value=self.openai_api_key)
        openai_model_var = tk.StringVar(value=self.openai_model)
        ollama_url_var = tk.StringVar(value=self.ollama_url)
        ollama_model_var = tk.StringVar(value=self.ollama_model)
        output_dir_var = tk.StringVar(value=self.output_directory)
        json_output_dir_var = tk.StringVar(value=self.json_output_directory)
        initiator_var = tk.StringVar(value=self.initiator)
        
        def show_provider_fields(event=None):
            for widget in config_frame.winfo_children():
                widget.destroy()
            provider = self.ai_provider_var.get()
            if provider == "openai":
                ttk.Label(config_frame, text="API Key:").pack(anchor=tk.W)
                ttk.Entry(config_frame, textvariable=openai_api_key_var, show="*", width=50).pack(fill=tk.X, pady=(0, 5))
                ttk.Label(config_frame, text="Model:").pack(anchor=tk.W)
                ttk.Entry(config_frame, textvariable=openai_model_var, width=50).pack(fill=tk.X, pady=(0, 5))
            elif provider == "ollama":
                ttk.Label(config_frame, text="Ollama URL:").pack(anchor=tk.W)
                ttk.Entry(config_frame, textvariable=ollama_url_var, width=50).pack(fill=tk.X, pady=(0, 5))
                ttk.Label(config_frame, text="Model:").pack(anchor=tk.W)
                ttk.Entry(config_frame, textvariable=ollama_model_var, width=50).pack(fill=tk.X, pady=(0, 5))
        
        ai_provider_combo.bind("<<ComboboxSelected>>", show_provider_fields)
        show_provider_fields()

        # Document template selection
        ttk.Label(scrollable_frame, text="Document Template auswählen:").pack(anchor=tk.W)
        
        def get_valid_templates():
            valid_files = []
            valid_names = []
            # Alle aktuellen Section-Platzhalter
            placeholders = [
                "system_scope",
                "architecture_tech_stack",
                "external_interfaces",
                "ci_cd",
                "testing_concept",
                "deployment_operation",
                "ux_ui"
            ]
            # Regex für alle Platzhalter (mit/ohne Leerzeichen, case-insensitive)
            placeholder_patterns = [re.compile(r"\{\{\s*" + ph + r"\s*\}\}", re.IGNORECASE) for ph in placeholders]
            for f in glob.glob("templates/*.docx") + glob.glob("templates/*.dotx"):
                try:
                    d = docx.Document(f)
                    found = False
                    # Check paragraphs
                    for p in d.paragraphs:
                        for pat in placeholder_patterns:
                            if pat.search(p.text):
                                found = True
                                print(f"[TEMPLATE-DEBUG] ACCEPTED: {f} (found in paragraph: {p.text})")
                                break
                        if found:
                            break
                    # Check tables if not found yet
                    if not found:
                        for table in d.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for pat in placeholder_patterns:
                                        if pat.search(cell.text):
                                            found = True
                                            print(f"[TEMPLATE-DEBUG] ACCEPTED: {f} (found in table cell: {cell.text})")
                                            break
                                    if found:
                                        break
                                if found:
                                    break
                            if found:
                                break
                    if found:
                        valid_files.append(f)
                        valid_names.append(os.path.basename(f))
                    else:
                        print(f"[TEMPLATE-DEBUG] REJECTED (no section placeholder): {f}")
                except Exception as e:
                    print(f"[TEMPLATE-DEBUG] ERROR: {f} - {e}")
            return valid_files, valid_names
        
        template_files, template_names = get_valid_templates()
        self.template_var = tk.StringVar(value=self.selected_template or (template_names[0] if template_names else ""))
        template_combo = ttk.Combobox(scrollable_frame, textvariable=self.template_var, values=template_names, state="readonly")
        template_combo.pack(fill=tk.X, pady=(0, 20))
        
        def refresh_templates():
            files, names = get_valid_templates()
            template_combo["values"] = names
            if names:
                self.template_var.set(names[0])
            else:
                self.template_var.set("")
            # Update the template_files and template_names in the enclosing scope
            nonlocal template_files, template_names
            template_files, template_names = files, names
        
        refresh_btn = ttk.Button(scrollable_frame, text="Templates aktualisieren", command=refresh_templates)
        refresh_btn.pack(pady=(0, 10))
        
        # Output directory setting
        ttk.Label(scrollable_frame, text="Ausgabeverzeichnis:").pack(anchor=tk.W, pady=(15,0))
        output_dir_frame = ttk.Frame(scrollable_frame)
        output_dir_frame.pack(fill=tk.X, pady=(0, 10))
        output_dir_entry = ttk.Entry(output_dir_frame, textvariable=output_dir_var, width=40)
        output_dir_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        def browse_output_dir():
            dir_path = filedialog.askdirectory(title="Ausgabeverzeichnis auswählen", initialdir=output_dir_var.get())
            if dir_path:
                output_dir_var.set(dir_path)
        
        browse_btn = ttk.Button(output_dir_frame, text="Durchsuchen", command=browse_output_dir)
        browse_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # JSON Output directory setting
        ttk.Label(scrollable_frame, text="JSON Ausgabeverzeichnis:").pack(anchor=tk.W, pady=(15,0))
        json_output_dir_frame = ttk.Frame(scrollable_frame)
        json_output_dir_frame.pack(fill=tk.X, pady=(0, 10))
        json_output_dir_entry = ttk.Entry(json_output_dir_frame, textvariable=json_output_dir_var, width=40)
        json_output_dir_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        def browse_json_output_dir():
            dir_path = filedialog.askdirectory(title="JSON Ausgabeverzeichnis auswählen", initialdir=json_output_dir_var.get())
            if dir_path:
                json_output_dir_var.set(dir_path)
        
        browse_json_btn = ttk.Button(json_output_dir_frame, text="Durchsuchen", command=browse_json_output_dir)
        browse_json_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # Akzeptanzschwellen-Slider
        ttk.Label(scrollable_frame, text="Akzeptanzschwelle für inhaltliche Übereinstimmung (0.3 = locker, 1.0 = sehr streng):").pack(anchor=tk.W, pady=(15,0))
        threshold_var = tk.DoubleVar(value=self.alignment_threshold)
        threshold_label = ttk.Label(scrollable_frame, text=f"Aktueller Wert: {threshold_var.get():.2f}")
        threshold_label.pack(anchor=tk.W)
        
        def on_slider_change(val):
            threshold_label.config(text=f"Aktueller Wert: {float(val):.2f}")
        
        threshold_slider = ttk.Scale(scrollable_frame, from_=0.3, to=1.0, orient=tk.HORIZONTAL, variable=threshold_var, command=on_slider_change)
        threshold_slider.pack(fill=tk.X, pady=(0, 20))
        
        # Initiator setting
        ttk.Label(scrollable_frame, text="Initiator (optional):").pack(anchor=tk.W, pady=(15,0))
        ttk.Label(scrollable_frame, text="Name/Abteilung des Veranlassers (wird als Suffix an Dateinamen angehängt):", font=("Arial", 8)).pack(anchor=tk.W)
        initiator_entry = ttk.Entry(scrollable_frame, textvariable=initiator_var, width=40)
        initiator_entry.pack(fill=tk.X, pady=(0, 20))

        def save_settings():
            # Save document template
            # Always use the latest template_files/names
            files, names = get_valid_templates()
            if self.template_var.get():
                idx = names.index(self.template_var.get()) if self.template_var.get() in names else -1
                if idx >= 0:
                    self.selected_template = files[idx]
                    # Set template in word generator
                    self.word_generator.set_template(self.selected_template)
                    if hasattr(self, 'logger') and self.logger:
                        self.logger.info("Template set in word generator: %s", self.selected_template)
            else:
                self.selected_template = None
            
            provider = self.ai_provider_var.get()
            if provider == "openai":
                self.openai_api_key = openai_api_key_var.get()
                self.openai_model = openai_model_var.get()
                os.environ["OPENAI_API_KEY"] = self.openai_api_key
                os.environ["OPENAI_MODEL"] = self.openai_model
            elif provider == "ollama":
                self.ollama_url = ollama_url_var.get()
                self.ollama_model = ollama_model_var.get()
                os.environ["OLLAMA_URL"] = self.ollama_url
                os.environ["OLLAMA_MODEL"] = self.ollama_model
            
            self.alignment_threshold = float(threshold_var.get())
            self.output_directory = output_dir_var.get()
            self.json_output_directory = json_output_dir_var.get()
            self.initiator = initiator_var.get()
            self.save_config()
            # Reinitialize services with new output directory
            self._reinitialize_services()
            settings_window.destroy()
        
        # Save button
        save_btn = ttk.Button(scrollable_frame, text="Speichern", command=save_settings)
        save_btn.pack(pady=(20, 0))
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Configure canvas scrolling
        settings_window.bind("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

    def load_from_file(self):
        """Load project data from JSON file"""
        file_path = filedialog.askopenfilename(
            title="Select Project Data File",
            filetypes=[("JSON files", "*.json"), ("Text files", "*.txt"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    
                    # Try to parse as JSON first
                    try:
                        import json
                        data = json.loads(content)
                        
                        # Extract fields from JSON object
                        project_name = data.get('name', '').strip()
                        upwork_link = data.get('link', '').strip()
                        description = data.get('description', '').strip()
                        
                        # Update GUI fields
                        if project_name:
                            self.project_name_var.set(project_name)
                        if upwork_link:
                            self.upwork_link_var.set(upwork_link)
                        if description:
                            self.description_text.delete(1.0, tk.END)
                            self.description_text.insert(1.0, description)
                        
                        # Show what was loaded
                        loaded_info = []
                        if project_name:
                            loaded_info.append(f"Name: {project_name}")
                        if upwork_link:
                            loaded_info.append(f"Link: {upwork_link}")
                        if description:
                            loaded_info.append(f"Description: {len(description)} chars")
                        
                        status_msg = f"Loaded JSON from {Path(file_path).name}: {', '.join(loaded_info)}"
                        self.status_var.set(status_msg)
                        
                    except json.JSONDecodeError:
                        # Fallback to old text parsing for backward compatibility
                        project_name = ""
                        upwork_link = ""
                        description = ""
                        
                        lines = content.split('\n')
                        current_section = None
                        
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                                
                            # Check for [project_name] style placeholders
                            if line.lower().startswith('[name]') or line.lower().startswith('[project_name]'):
                                # Extract content after the placeholder
                                if line.lower().startswith('[name]'):
                                    project_name = line[6:].strip()
                                else:
                                    project_name = line[15:].strip()
                            elif line.lower().startswith('[link]') or line.lower().startswith('[upwork_link]'):
                                # Extract content after the placeholder
                                if line.lower().startswith('[link]'):
                                    upwork_link = line[7:].strip()
                                else:
                                    upwork_link = line[13:].strip()
                            elif line.lower().startswith('[description]') or line.lower().startswith('[summary]'):
                                current_section = 'description'
                                if line.lower().startswith('[description]'):
                                    description = line[13:].strip()
                                else:
                                    description = line[9:].strip()
                            elif current_section == 'description':
                                # Continue reading description
                                description += '\n' + line
                            else:
                                # If no section headers found, treat as description
                                if not project_name and not upwork_link:
                                    description += line + '\n'
                        
                        # Update GUI fields
                        if project_name:
                            self.project_name_var.set(project_name)
                        if upwork_link:
                            self.upwork_link_var.set(upwork_link)
                        if description:
                            self.description_text.delete(1.0, tk.END)
                            self.description_text.insert(1.0, description.strip())
                        
                        # Show what was loaded
                        loaded_info = []
                        if project_name:
                            loaded_info.append(f"Name: {project_name}")
                        if upwork_link:
                            loaded_info.append(f"Link: {upwork_link}")
                        if description:
                            loaded_info.append(f"Description: {len(description)} chars")
                        
                        status_msg = f"Loaded text from {Path(file_path).name}: {', '.join(loaded_info)}"
                        self.status_var.set(status_msg)
                    
            except Exception as e:
                messagebox.showerror("Error", f"Could not load file: {str(e)}")
    
    def generate_json_from_specification(self):
        """Generate JSON file from current specification in the UI"""
        # Get project data from UI
        project_name = self.project_name_var.get().strip()
        if not project_name or project_name == "Enter project name here...":
            messagebox.showwarning("Warning", "Please enter a project name.")
            return
            
        upwork_link = self.upwork_link_var.get().strip()
        description = self.description_text.get("1.0", tk.END).strip()
        if not description:
            messagebox.showwarning("Warning", "Please enter a project description.")
            return
        
        # Create JSON data structure
        json_data = {
            "name": project_name,
            "link": upwork_link,
            "description": description
        }
        
        # Create output directory if it doesn't exist
        json_output_dir = Path(self.json_output_directory)
        json_output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename using project name with better sanitization
        import re
        # Remove or replace invalid filename characters
        safe_project_name = re.sub(r'[<>:"/\\|?*]', '_', project_name)
        # Replace multiple spaces/underscores with single underscore
        safe_project_name = re.sub(r'[_\s]+', '_', safe_project_name)
        # Remove leading/trailing underscores and spaces
        safe_project_name = safe_project_name.strip('_ ')
        # Ensure the filename is not empty
        if not safe_project_name:
            safe_project_name = "unnamed_project"
        # Limit length to avoid filesystem issues
        if len(safe_project_name) > 100:
            safe_project_name = safe_project_name[:100]
        
        filename = f"{safe_project_name}.json"
        json_path = json_output_dir / filename
        
        try:
            # Write JSON file
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False)
            
            # Show success message
            message = f"JSON specification generated successfully!\n\nFile created:\n• {json_path}"
            messagebox.showinfo("Success", message)
            
            # Update status
            self.status_var.set(f"JSON file created: {filename}")
            
        except Exception as e:
            error_message = f"Error generating JSON file: {str(e)}"
            messagebox.showerror("Error", error_message)
            self.status_var.set("JSON generation failed")
    
    def generate_concept(self):
        """Start the concept generation process"""
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Starting concept generation from GUI")
        
        # Get project name from input
        project_name = self.project_name_var.get().strip()
        if not project_name or project_name == "Enter project name here...":
            if hasattr(self, 'logger') and self.logger:
                self.logger.warning("No project name provided")
            messagebox.showwarning("Warning", "Please enter a project name.")
            return
        upwork_link = self.upwork_link_var.get().strip()
        # upwork_link ist jetzt wirklich optional, keine Prüfung auf Platzhalter mehr
        description = self.description_text.get("1.0", tk.END).strip()
        if not description:
            if hasattr(self, 'logger') and self.logger:
                self.logger.warning("No project description provided")
            messagebox.showwarning("Warning", "Please enter a project description.")
            return
            
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Project description length: %d characters", len(description))
        
        # Disable generate button and enable cancel button
        self.generate_btn.config(state="disabled")
        self.cancel_btn.config(state="normal")
        self.cancel_requested = False
        
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Starting generation thread")
        # Start generation in a separate thread
        thread = threading.Thread(target=self._generate_concept_thread, args=(description, project_name, upwork_link))
        thread.daemon = True
        thread.start()

    def cancel_generation(self):
        """Cancel the current generation process"""
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Cancellation requested by user")
        self.cancel_requested = True
        self.cancel_btn.config(state="disabled")
        self.generate_btn.config(state="normal")
        self.progress_var.set("Cancelled")
        self.status_var.set("Generation cancelled")

    def _generate_concept_thread(self, description, project_name, upwork_link):
        """Generate concept in a separate thread"""
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Generation thread started")
        try:
            self.setup_logger()
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Logger setup completed")
            
            # Update GUI
            self.root.after(0, lambda: self.progress_var.set("Setting up..."))
            self.root.after(0, lambda: self.status_var.set("Initializing..."))
            
            # Set AI provider settings
            provider = self.ai_provider_var.get()
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Using AI provider: %s", provider)
            
            if provider == "openai":
                os.environ["OPENAI_API_KEY"] = self.openai_api_key
                os.environ["OPENAI_MODEL"] = self.openai_model
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("OpenAI settings configured")
            elif provider == "ollama":
                os.environ["OLLAMA_URL"] = self.ollama_url
                os.environ["OLLAMA_MODEL"] = self.ollama_model
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("Ollama settings configured")
                
            # Set alignment threshold
            self.ai_service.alignment_threshold = self.alignment_threshold
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Alignment threshold set to: %.2f", self.alignment_threshold)
            
            # Load proposal context if available
            proposal_context = self._load_full_proposal_context()
            if proposal_context:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("Loaded proposal context, length: %d characters", len(proposal_context))
            else:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("No proposal context found")
            
            # Update GUI
            self.root.after(0, lambda: self.progress_var.set("Generating concept..."))
            self.root.after(0, lambda: self.status_var.set("Generating technical concept..."))
            
            # Generate concept
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Starting AI concept generation")
            concept = self.ai_service.generate_technical_concept_sections(
                description, 
                provider=provider, 
                proposal_context=proposal_context,
                cancel_callback=lambda: self.cancel_requested
            )
            
            if self.cancel_requested:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("Generation cancelled during AI processing")
                return
                
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("AI concept generation completed")
                self.logger.info("Generated sections: %s", list(concept.get('sections', {}).keys()))
            

                
            # Use manually entered project name
            if hasattr(self, 'logger') and self.logger:
                self.logger.info(f"Using manually entered project name: {project_name}")

            # Word-Dokument erzeugen, Projektname, Datum und Initiator als Parameter übergeben
            docx_path = self.word_generator.create_document(
                concept, 
                project_name=project_name, 
                initiator=self.initiator,
                upwork_link=upwork_link,
                description=description
            )
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Word document created: %s", docx_path)
            # Öffne das erzeugte docx automatisch
            try:
                import platform
                import subprocess
                if platform.system() == "Windows":
                    os.startfile(docx_path)
                elif platform.system() == "Darwin":
                    subprocess.Popen(["open", docx_path])
                else:
                    subprocess.Popen(["xdg-open", docx_path])
            except Exception as e:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.error(f"Could not open docx automatically: {e}")
            
            if self.cancel_requested:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("Generation cancelled during Word document creation")
                return
                
            # Update GUI
            self.root.after(0, lambda: self.progress_var.set("Completed"))
            self.root.after(0, lambda: self.status_var.set("Generation completed successfully"))
            
            # Show completion message (nur docx)
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Generation process completed successfully")
            message = f"Technical concept generated successfully!\n\nFiles created:\n• {docx_path}"
            message += f"\n\nLog file: {self.logfile_path}"
            self.root.after(0, lambda: messagebox.showinfo("Success", message))
            
        except Exception as e:
            if hasattr(self, 'logger') and self.logger:
                self.logger.error("Generation failed: %s", str(e), exc_info=True)
            error_message = f"Error during generation: {str(e)}\n\nCheck the log file for details: {self.logfile_path}"
            self.root.after(0, lambda: messagebox.showerror("Error", error_message))
            self.root.after(0, lambda: self.progress_var.set("Error"))
            self.root.after(0, lambda: self.status_var.set("Generation failed"))
        finally:
            # Re-enable generate button and disable cancel button
            self.root.after(0, lambda: self.generate_btn.config(state="normal"))
            self.root.after(0, lambda: self.cancel_btn.config(state="disabled"))
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("Generation thread finished")

    def _load_full_proposal_context(self):
        """Load the full proposal context from the proposal file"""
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Loading proposal context")
        try:
            proposal_file = "proposal.txt"
            if os.path.exists(proposal_file):
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("Found proposal file: %s", proposal_file)
                with open(proposal_file, "r", encoding="utf-8") as f:
                    content = f.read().strip()
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("Loaded proposal content, length: %d characters", len(content))
                return content
            else:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.info("No proposal file found")
                return ""
        except Exception as e:
            if hasattr(self, 'logger') and self.logger:
                self.logger.error("Error loading proposal context: %s", str(e))
            return ""

    def _convert_to_pdf(self, docx_path):
        """Convert Word document to PDF"""
        if hasattr(self, 'logger') and self.logger:
            self.logger.info("Converting Word document to PDF: %s", docx_path)
        
        try:
            # Create PDF output directory
            pdf_output_dir = Path(self.output_directory) / "pdf"
            pdf_output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate PDF filename
            docx_name = Path(docx_path).stem
            pdf_path = pdf_output_dir / f"{docx_name}.pdf"
            
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("PDF output path: %s", pdf_path)
            
            # Initialize COM for Windows
            import platform
            if platform.system() == "Windows":
                try:
                    import pythoncom
                    pythoncom.CoInitialize()
                    if hasattr(self, 'logger') and self.logger:
                        self.logger.info("COM initialized for Windows")
                except ImportError:
                    if hasattr(self, 'logger') and self.logger:
                        self.logger.warning("pythoncom not available, trying without COM initialization")
                except Exception as e:
                    if hasattr(self, 'logger') and self.logger:
                        self.logger.warning("COM initialization failed: %s", str(e))
            # PDF conversion is optional and skipped if docx2pdf is not installed
            if hasattr(self, 'logger') and self.logger:
                self.logger.info("PDF conversion skipped (docx2pdf not installed)")
            return None
        except Exception as e:
            if hasattr(self, 'logger') and self.logger:
                self.logger.error("PDF conversion failed: %s", str(e))
            raise

    def show_last_log(self):
        if self.logfile_path and os.path.exists(self.logfile_path):
            with open(self.logfile_path, "r", encoding="utf-8") as f:
                log_content = f.read()
            log_win = tk.Toplevel(self.root)
            log_win.title(f"Log: {os.path.basename(self.logfile_path)}")
            text = scrolledtext.ScrolledText(log_win, width=100, height=30)
            text.pack(fill=tk.BOTH, expand=True)
            text.insert(tk.END, log_content)
            text.config(state="disabled")
        else:
            messagebox.showinfo("Info", "Kein Logfile gefunden.")

    def show_preview(self, png_path):
        preview_win = tk.Toplevel(self.root)
        preview_win.title("Document Preview")
        img = Image.open(png_path)
        tk_img = ImageTk.PhotoImage(img)
        label = tk.Label(preview_win, image=tk_img)
        label.image = tk_img  # Prevent garbage collection  # type: ignore[attr-defined]
        label.pack()
        preview_win.geometry(f"{img.width}x{img.height}")

    def log_message(self, message):
        print(message)
        if self.log_widget:
            self.log_widget.config(state="normal")
            self.log_widget.insert(tk.END, message + "\n")
            self.log_widget.see(tk.END)
            self.log_widget.config(state="disabled")

    def toggle_log(self):
        if self.log_visible:
            if self.log_label:
                self.log_label.grid_remove()
            if self.log_widget:
                self.log_widget.grid_remove()
            self.toggle_log_btn.config(text="Log anzeigen")
            self.log_visible = False
        else:
            if self.log_label:
                self.log_label.grid(row=7, column=0, sticky=tk.W, pady=(10, 0))
            if self.log_widget:
                self.log_widget.grid(row=8, column=0, columnspan=2, sticky="ew", pady=(0, 10))
            self.toggle_log_btn.config(text="Log ausblenden")
            self.log_visible = True

def main():
    root = tk.Tk()
    app = ZetaProposerGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main() 