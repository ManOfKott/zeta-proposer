from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any, List, Optional
from pathlib import Path
import os
import re
import logging
from datetime import datetime
import tkinter.messagebox as messagebox


class WordDocumentGenerator:
    def __init__(self, output_directory="output"):
        self.output_dir = Path(output_directory)
        self.output_dir.mkdir(exist_ok=True)
        self.template_path = None  # Path to selected template
        self.logger = logging.getLogger(__name__)
        self.logger.info("WordDocumentGenerator initialized")
        self.logger.info("Output directory: %s", self.output_dir)
    
    def set_template(self, template_path: Optional[str]):
        self.logger.info("Setting template: %s", template_path)
        self.template_path = template_path
    
    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting from text"""
        self.logger.debug("Cleaning markdown from text")
        self.logger.debug("Input text length: %d", len(text) if text else 0)
        
        if not text:
            self.logger.debug("No text to clean")
            return text
        
        # Remove markdown headers (# ## ### etc.)
        text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
        
        # Remove bold formatting (**text** or __text__)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        
        # Remove italic formatting (*text* or _text_)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)
        
        # Remove code formatting (`text`)
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        # Remove bullet points and dashes at the beginning of lines
        text = re.sub(r'^[\s]*[-*+]\s+', '', text, flags=re.MULTILINE)
        
        # Clean up extra whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        self.logger.debug("Cleaned text length: %d", len(text))
        return text
    
    def _remove_mermaid_blocks(self, text: str) -> str:
        self.logger.debug("Removing mermaid blocks from text")
        self.logger.debug("Input text length: %d", len(text) if text else 0)
        
        if not text:
            return text
        # Entfernt ```mermaid ... ``` Blöcke
        text = re.sub(r'```mermaid[\s\S]+?```', '', text, flags=re.IGNORECASE)
        # Entfernt Zeilen, die mit 'mermaid' oder 'graph' beginnen
        lines = text.splitlines()
        cleaned_lines = [line for line in lines if not line.strip().lower().startswith(('mermaid', 'graph'))]
        result = '\n'.join(cleaned_lines)
        
        self.logger.debug("Text after mermaid removal length: %d", len(result))
        return result
    
    def _remove_section_heading(self, content: str, title: str) -> str:
        """Remove the section heading/title from the AI content if present."""
        self.logger.debug("Removing section heading for title: %s", title)
        self.logger.debug("Content length: %d", len(content) if content else 0)
        
        if not content:
            return ''
        lines = content.strip().split('\n')
        # Remove first line if it matches the title (with or without numbering)
        if lines and (title.lower() in lines[0].lower() or lines[0].strip().startswith(tuple(str(i) for i in range(1, 10)))):
            result = '\n'.join(lines[1:]).lstrip()
            self.logger.debug("Removed heading, new length: %d", len(result))
            return result
        self.logger.debug("No heading removed, content unchanged")
        return content

    def _insert_single_section(self, doc, insert_after, key, title, ai_sections, heading_style, normal_style):
        """Replace the placeholder in the paragraph/cell directly with content, keeping document order stable."""
        self.logger.info("Inserting single section: %s", key)
        self.logger.info("Section title: %s", title)
        
        # Hole Fließtext aus neuer Struktur
        value = ai_sections.get(key)
        if isinstance(value, dict):
            content = value.get('text', '')
        else:
            content = value
        if not content:
            for k, v in ai_sections.items():
                if title.lower() in k.replace('_', ' ').lower():
                    if isinstance(v, dict):
                        content = v.get('text', '')
                    else:
                        content = v
                    break
        self.logger.info("Section %s: value length = %d", key, len(content) if content else 0)
        self.logger.debug("Section %s: value preview = %s", key, content[:100] if content else 'None')
        
        # Remove heading/title from AI content
        content = self._remove_section_heading(content, title)
        content = self._remove_mermaid_blocks(content)
        content = self._remove_dot_blocks(content)
        
        if content and not content.startswith("Section") and not content.startswith("Error"):
            content = self._clean_markdown(content)
            self.logger.info("Section %s: cleaned content length = %d", key, len(content))
            self.logger.debug("Section %s: content preview = %s", key, content[:200])
        else:
            content = f"(No information available for {title})"
            self.logger.info("Section %s: using fallback content", key)
            
        # Paragraph case
        if insert_after is not None and hasattr(insert_after, 'text') and hasattr(insert_after, '_element'):
            self.logger.info("Inserting content in paragraph/cell for section: %s", key)
            # Ersetze nur den Text des Platzhalter-Absatzes
            insert_after.text = content
            self.logger.info("Replaced placeholder for section %s with content in paragraph/cell", key)
            return
        # Table cell case
        elif hasattr(insert_after, 'text') and hasattr(insert_after, 'add_paragraph'):
            self.logger.info("Inserting content in table cell for section: %s", key)
            insert_after.text = content  # Platzhaltertext ersetzen
            self.logger.info("Replaced placeholder for section %s with content in table cell", key)
            return
        else:
            self.logger.warning("Unknown insert_after type for section %s: %s", key, type(insert_after))
            return

    def _escape_xml_text(self, text: str) -> str:
        """Escape special characters for XML"""
        if not text:
            return text
        # Replace special characters that can cause XML issues
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&apos;')
        return text
    
    def _ai_shorten_project_name(self, project_name: str, max_length: int, max_retries: int = 3) -> str:
        """KI-basierte Kürzung des Projektnamens mit Retry-Logik"""
        # Bereinige den Namen zuerst
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', project_name)
        safe_name = re.sub(r'\s+', '_', safe_name)
        
        # Wenn der Name bereits kurz genug ist, verwende ihn direkt
        if len(safe_name) <= max_length:
            return safe_name
        
        # KI-basierte Kürzung mit Retry-Logik
        for attempt in range(max_retries):
            try:
                # Verwende OpenAI für intelligente Kürzung
                import os
                from openai import OpenAI
                
                api_key = os.getenv("OPENAI_API_KEY")
                if api_key:
                    client = OpenAI(api_key=api_key)
                    
                    prompt = f"""Kürze den folgenden Projektnamen intelligent auf maximal {max_length} Zeichen.
                    
                    Regeln:
                    - Behalte die wichtigsten Wörter bei
                    - Entferne unwichtige Wörter wie "Projekt", "System", "Anwendung" wenn nötig
                    - Verwende Abkürzungen wo sinnvoll (z.B. "Management" -> "Mgmt", "Technical" -> "Tech")
                    - Ersetze Leerzeichen durch Unterstriche
                    - Entferne Sonderzeichen außer Unterstrichen
                    - Verwende KEINE Punkte (...) am Ende
                    - Kürze so, dass der Name vollständig und verständlich bleibt
                    - WICHTIG: Der Name darf maximal {max_length} Zeichen haben
                    
                    Originaler Name: {project_name}
                    
                    Gib nur den gekürzten Namen zurück, ohne Erklärung."""
                    
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=50,
                        temperature=0.3
                    )
                    
                    content = response.choices[0].message.content
                    if content is None:
                        raise Exception("OpenAI returned empty response")
                    shortened_name = content.strip()
                    
                    # Bereinige das Ergebnis
                    shortened_name = re.sub(r'[<>:"/\\|?*]', '_', shortened_name)
                    shortened_name = re.sub(r'\s+', '_', shortened_name)
                    
                    # Prüfe Länge und retry falls nötig
                    if len(shortened_name) <= max_length:
                        self.logger.info(f"KI-Kürzung erfolgreich (Versuch {attempt + 1}): '{project_name}' -> '{shortened_name}'")
                        return shortened_name
                    else:
                        self.logger.warning(f"KI-Kürzung zu lang (Versuch {attempt + 1}): {len(shortened_name)} > {max_length} Zeichen")
                        if attempt < max_retries - 1:
                            continue  # Retry
                        else:
                            # Letzter Versuch: Manuelle Kürzung
                            shortened_name = shortened_name[:max_length]
                            self.logger.info(f"Manuelle Kürzung nach {max_retries} KI-Versuchen: '{project_name}' -> '{shortened_name}'")
                            return shortened_name
                    
            except Exception as e:
                self.logger.warning(f"KI-Kürzung fehlgeschlagen (Versuch {attempt + 1}): {e}")
                if attempt < max_retries - 1:
                    continue  # Retry
                else:
                    # Letzter Versuch: Manuelle Kürzung
                    shortened_name = safe_name[:max_length]
                    self.logger.info(f"Manuelle Kürzung nach {max_retries} fehlgeschlagenen KI-Versuchen: '{project_name}' -> '{shortened_name}'")
                    return shortened_name
        
        # Fallback: Manuelle Kürzung
        shortened_name = safe_name[:max_length]
        self.logger.info(f"Manuelle Kürzung: '{project_name}' -> '{shortened_name}'")
        return shortened_name

    def _replace_placeholders_in_xml(self, doc, replacements: dict):
        """Replace placeholders in XML elements (shapes, textboxes, etc.)"""
        self.logger.info("Replacing placeholders in XML elements")
        for element in doc.element.xpath('//w:t'):
            if element.text:
                for placeholder, replacement in replacements.items():
                    if placeholder in element.text:
                        # Kein XML-Escaping mehr, sondern roher Text
                        element.text = element.text.replace(placeholder, replacement)
                        self.logger.info("Replaced %s in XML element", placeholder)

    def create_document(self, concept: Dict[str, Any], project_name: Optional[str] = None, initiator: Optional[str] = None, upwork_link: Optional[str] = None, description: Optional[str] = None) -> str:
        """Create a Word document from the concept data using template replacement. Speichert alles im Projektordner mit Versionierung."""
        self.logger.info("Starting document creation")
        
        # --- Namensgenerierung: KI nur wenn Limit überschritten ---
        max_name_len = 25
        # Bereinige den ursprünglichen Namen
        original_safe_name = re.sub(r'[<>:"/\\|?*]', '_', project_name or "Technical_Concept")
        original_safe_name = re.sub(r'\s+', '_', original_safe_name)
        
        # Verwende KI nur wenn der ursprüngliche Name zu lang ist
        if len(original_safe_name) > max_name_len:
            safe_project_name = self._ai_shorten_project_name(project_name or "Technical_Concept", max_name_len)
            self.logger.info(f"Projektname überschreitet Limit ({len(original_safe_name)} > {max_name_len}), verwende KI-Kürzung")
        else:
            safe_project_name = original_safe_name
            self.logger.info(f"Projektname innerhalb des Limits ({len(original_safe_name)} <= {max_name_len}), verwende Original")
        
        # --- Veranlasser als Präfix hinzufügen ---
        safe_initiator = re.sub(r'[<>:"/\\|?*]', '_', initiator or "")
        safe_initiator = re.sub(r'\s+', '_', safe_initiator)
        
        # Ordnername: Veranlasser_Projektname (KI-generiert)
        if safe_initiator:
            base_folder_name = f"{safe_initiator}_{safe_project_name}"
        else:
            base_folder_name = safe_project_name
            
        # Versionierung
        version = 1
        project_folder = self.output_dir / base_folder_name
        while project_folder.exists():
            version += 1
            project_folder = self.output_dir / f"{base_folder_name}_v{version}"
        
        # Stelle sicher, dass der Ordner erstellt wird
        try:
            project_folder.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"Projektordner erstellt: {project_folder}")
        except Exception as e:
            self.logger.error(f"Fehler beim Erstellen des Projektordners: {e}")
            # Fallback: Verwende einen einfacheren Pfad
            fallback_folder = Path("output") / f"fallback_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            fallback_folder.mkdir(parents=True, exist_ok=True)
            project_folder = fallback_folder
            self.logger.info(f"Fallback-Ordner verwendet: {project_folder}")

        # --- DOCX-Datei erzeugen (wie bisher, aber im Projektordner) ---
        # (Kopiere bisherigen Code für Template/Platzhalterersetzung, aber speichere in project_folder)
        # Extract AI sections
        ai_sections = concept.get('sections', {})
        self.logger.info("Concept sections count: %d", len(ai_sections))
        
        # Use template if available
        if self.template_path and os.path.exists(self.template_path):
            self.logger.info("Template path: %s", self.template_path)
            self.logger.info("Template exists: %s", os.path.exists(self.template_path))
            self.logger.info("Using existing template: %s", self.template_path)
            
            # Load template
            doc = Document(self.template_path)
            
            # Extract sections for logging
            self.logger.info("Extracted ai_sections:")
            for key, value in ai_sections.items():
                if isinstance(value, dict):
                    content = value.get('text', '')
                else:
                    content = value
                self.logger.info("  %s: %s", key, content[:100] if content else 'None')
            
            # Check if template has placeholders
            has_placeholders = False
            found_placeholders = []
            
            # Check paragraphs for placeholders
            self.logger.info("Checking paragraphs for placeholders")
            for paragraph in doc.paragraphs:
                for key in ai_sections.keys():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        has_placeholders = True
                        found_placeholders.append(key)
                        self.logger.info("Found placeholder %s in paragraph: %s...", placeholder, paragraph.text[:50])
                        break
            
            # Check tables for placeholders
            self.logger.info("Checking tables for placeholders")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key in ai_sections.keys():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                has_placeholders = True
                                if key not in found_placeholders:
                                    found_placeholders.append(key)
                                self.logger.info("Found placeholder %s in table cell: %s...", placeholder, cell.text[:50])
                                break
            
            self.logger.info("Template has placeholders: %s", has_placeholders)
            
            if has_placeholders:
                self.logger.info("Found placeholders: %s", found_placeholders)
                self.logger.info("Replacing placeholders in existing template")
                
                # Process each placeholder
                for key in found_placeholders:
                    self.logger.info("Processing placeholder for section: %s", key)
                    placeholder = f"{{{{{key}}}}}"
                    
                    # Find and replace in paragraphs
                    for paragraph in doc.paragraphs:
                        if placeholder in paragraph.text:
                            self.logger.info("Found placeholder %s in paragraph. Inserting content.", placeholder)
                            self._insert_single_section(doc, paragraph, key, key.replace('_', ' ').title(), ai_sections, None, None)
                            break
                    
                    # Find and replace in table cells
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if placeholder in cell.text:
                                    self.logger.info("Found placeholder %s in table cell. Inserting content.", placeholder)
                                    self._insert_single_section(doc, cell, key, key.replace('_', ' ').title(), ai_sections, None, None)
                                    break
                            else:
                                continue
                            break
                        else:
                            continue
                        break
                
                # Replace date and project name placeholders in paragraphs and table cells
                timestamp = datetime.now().strftime("%Y-%m-%d")
                safe_project_name = re.sub(r'[<>:"/\\|?*]', '_', project_name or "Technical_Concept")
                safe_project_name = re.sub(r'\s+', '_', safe_project_name)
                
                # Replace in paragraphs
                for paragraph in doc.paragraphs:
                    if "{{date}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{date}}", timestamp)
                        self.logger.info("Replaced {{date}} in paragraph")
                    if "{{project_name}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{project_name}}", project_name or "Technical Concept")
                        self.logger.info("Replaced {{project_name}} in paragraph")
                    if "{{project_name_safe}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{project_name_safe}}", safe_project_name)
                        self.logger.info("Replaced {{project_name_safe}} in paragraph")
                
                # Replace in table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "{{date}}" in cell.text:
                                cell.text = cell.text.replace("{{date}}", timestamp)
                                self.logger.info("Replaced {{date}} in table cell")
                            if "{{project_name}}" in cell.text:
                                cell.text = cell.text.replace("{{project_name}}", project_name or "Technical Concept")
                                self.logger.info("Replaced {{project_name}} in table cell")
                            if "{{project_name_safe}}" in cell.text:
                                cell.text = cell.text.replace("{{project_name_safe}}", safe_project_name)
                                self.logger.info("Replaced {{project_name_safe}} in table cell")
                
                # Replace placeholders in XML elements (shapes, textboxes)
                replacements = {}
                for key, value in ai_sections.items():
                    if isinstance(value, dict):
                        content = value.get('text', '')
                    else:
                        content = value
                    if content:
                        content = self._clean_markdown(content)
                        replacements[f"{{{{{key}}}}}"] = content
                
                # Add date and project name placeholders
                timestamp = datetime.now().strftime("%Y-%m-%d")
                safe_project_name = re.sub(r'[<>:"/\\|?*]', '_', project_name or "Technical_Concept")
                safe_project_name = re.sub(r'\s+', '_', safe_project_name)
                
                replacements["{{date}}"] = timestamp
                replacements["{{project_name}}"] = project_name or "Technical Concept"
                replacements["{{project_name_safe}}"] = safe_project_name
                
                self._replace_placeholders_in_xml(doc, replacements)
                
                # Generate filename with retry logic (verwende gleichen KI-generierten Namen)
                # Verwende den bereits generierten safe_project_name für den Dateinamen
                
                # Add initiator suffix if provided
                initiator_suffix = f"_{initiator}" if initiator else ""
                
                # Try to save with versioning logic
                # Verwende den gleichen KI-generierten Namen für den Dateinamen
                base_filename = f"{safe_project_name}{initiator_suffix}.docx"
                version = 1
                
                while True:
                    try:
                        if version == 1:
                            filename = base_filename
                        else:
                            # Insert version before the extension
                            name_part = base_filename[:-5]  # Remove .docx
                            filename = f"{name_part}_v{version}.docx"
                        
                        docx_path = project_folder / filename
                        
                        # Ensure output directory exists
                        project_folder.mkdir(parents=True, exist_ok=True)
                        
                        self.logger.info("Saving document to: %s (version %d)", docx_path, version)
                        
                        # Use absolute path and handle long paths
                        abs_output_path = docx_path.resolve()
                        self.logger.info("Absolute path: %s", abs_output_path)
                        
                        doc.save(str(abs_output_path))
                        self.logger.info("Document saved successfully: %s", abs_output_path)
                        self.logger.info("Word document created: %s", abs_output_path)
                        break
                        
                    except PermissionError as e:
                        self.logger.warning("Permission error on version %d: %s", version, str(e))
                        version += 1
                        if version > 100:  # Prevent infinite loop
                            # Try to save to a simpler path as fallback
                            try:
                                fallback_path = Path("output") / f"fallback_{timestamp}_v{version}.docx"
                                self.logger.info("Trying fallback path: %s", fallback_path)
                                doc.save(str(fallback_path))
                                self.logger.info("Document saved to fallback path: %s", fallback_path)
                                return str(fallback_path)
                            except Exception as fallback_e:
                                self.logger.error("Fallback save also failed: %s", str(fallback_e))
                                raise e
                        
                    except Exception as e:
                        self.logger.error("Error saving document: %s", str(e))
                        raise e
                
                docx_path = project_folder / filename
            else:
                self.logger.warning("No placeholders found in template, creating new document")
                # Fallback to creating new document
                docx_path = self._create_new_document(concept, project_name, initiator)
        else:
            self.logger.info("No template available, creating new document")
            docx_path = self._create_new_document(concept, project_name, initiator)

        # --- TXT-Summary speichern ---
        try:
            summary_txt = f"Project Name: {project_name}\nDescription: {description}\nUpwork Link: {upwork_link or ''}\n"
            summary_path = project_folder / "summary.txt"
            with open(summary_path, "w", encoding="utf-8") as f:
                f.write(summary_txt)
            self.logger.info(f"Summary gespeichert: {summary_path}")
        except Exception as e:
            self.logger.error(f"Fehler beim Speichern der Summary: {e}")

        # --- Upwork-Link als .url-Datei speichern (nur wenn Link vorhanden) ---
        if upwork_link:
            try:
                url_path = project_folder / "upwork_link.url"
                with open(url_path, "w", encoding="utf-8") as f:
                    f.write(f"[InternetShortcut]\nURL={upwork_link}\n")
                self.logger.info(f"Upwork-Link gespeichert: {url_path}")
            except Exception as e:
                self.logger.error(f"Fehler beim Speichern der Upwork-Link-Datei: {e}")

        # --- Pfadlängen-Prüfung ---
        max_path_length = 240
        if len(str(docx_path)) > max_path_length:
            warn_msg = f"Der Pfad der Word-Datei ist zu lang (>{max_path_length} Zeichen)!\nBitte wähle einen kürzeren Projektnamen oder ein anderes Zielverzeichnis.\n\nPfad:\n{docx_path}"
            self.logger.error(warn_msg)
            try:
                messagebox.showerror("Pfad zu lang", warn_msg)
            except Exception:
                pass  # Falls kein GUI-Kontext vorhanden ist
            raise OSError(warn_msg)

        return str(docx_path)

    def _create_new_document(self, concept: Dict[str, Any], project_name: Optional[str] = None, initiator: Optional[str] = None) -> str:
        """Create a new document from scratch (fallback method)"""
        self.logger.info("Creating new document from scratch")
        
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add title page
        self._add_title_page(doc, concept)
        
        # Add table of contents
        self._add_table_of_contents(doc)
        
        # Add project overview
        self._add_project_overview(doc, concept)
        
        # Add system architecture
        self._add_system_architecture(doc, concept)
        
        # Add technical requirements
        self._add_technical_requirements(doc, concept)
        
        # Add implementation approach
        self._add_implementation_approach(doc, concept)
        
        # Add metadata
        self._add_metadata(doc, concept)
        
        # Generate filename with versioning (KI nur wenn Limit überschritten)
        max_name_len = 25
        # Bereinige den ursprünglichen Namen
        original_safe_name = re.sub(r'[<>:"/\\|?*]', '_', project_name or "Technical_Concept")
        original_safe_name = re.sub(r'\s+', '_', original_safe_name)
        
        # Verwende KI nur wenn der ursprüngliche Name zu lang ist
        if len(original_safe_name) > max_name_len:
            safe_project_name = self._ai_shorten_project_name(project_name or "Technical_Concept", max_name_len)
        else:
            safe_project_name = original_safe_name
        
        # Add initiator suffix if provided
        initiator_suffix = f"_{initiator}" if initiator else ""
        
        base_filename = f"{safe_project_name}{initiator_suffix}.docx"
        version = 1
        
        while True:
            try:
                if version == 1:
                    filename = base_filename
                else:
                    # Insert version before the extension
                    name_part = base_filename[:-5]  # Remove .docx
                    filename = f"{name_part}_v{version}.docx"
                
                output_path = self.output_dir / filename
                
                # Ensure output directory exists
                self.output_dir.mkdir(parents=True, exist_ok=True)
                
                self.logger.info("Saving new document to: %s (version %d)", output_path, version)
                
                # Use absolute path
                abs_output_path = output_path.resolve()
                doc.save(str(abs_output_path))
                self.logger.info("New document created: %s", abs_output_path)
                break
                
            except PermissionError as e:
                self.logger.warning("Permission error on version %d: %s", version, str(e))
                version += 1
                if version > 100:  # Prevent infinite loop
                    raise e
        
        return str(output_path)

    def _setup_document_styles(self, doc):
        """Setup document styles"""
        # Title style
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading style
        heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Normal style
        normal_style = doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)

    def _add_title_page(self, doc, concept: Dict[str, Any]):
        """Add title page"""
        title = doc.add_paragraph("Technical Concept", style='CustomTitle')
        doc.add_paragraph()  # Spacing
        
        # Add project description if available
        if 'description' in concept:
            desc_para = doc.add_paragraph(concept['description'], style='CustomNormal')
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()

    def _add_table_of_contents(self, doc):
        """Add table of contents"""
        toc_heading = doc.add_paragraph("Table of Contents", style='CustomHeading')
        doc.add_paragraph("1. Project Overview", style='CustomNormal')
        doc.add_paragraph("2. System Architecture", style='CustomNormal')
        doc.add_paragraph("3. Technical Requirements", style='CustomNormal')
        doc.add_paragraph("4. Implementation Approach", style='CustomNormal')
        doc.add_page_break()

    def _add_project_overview(self, doc, concept: Dict[str, Any]):
        """Add project overview section"""
        doc.add_paragraph("1. Project Overview", style='CustomHeading')
        doc.add_paragraph("This section provides an overview of the technical concept and project scope.", style='CustomNormal')
        doc.add_paragraph()

    def _add_system_architecture(self, doc, concept: Dict[str, Any]):
        """Add system architecture section"""
        doc.add_paragraph("2. System Architecture", style='CustomHeading')
        doc.add_paragraph("This section describes the system architecture and technical approach.", style='CustomNormal')
        doc.add_paragraph()

    def _add_technical_requirements(self, doc, concept: Dict[str, Any]):
        """Add technical requirements section"""
        doc.add_paragraph("3. Technical Requirements", style='CustomHeading')
        doc.add_paragraph("This section outlines the technical requirements and specifications.", style='CustomNormal')
        doc.add_paragraph()

    def _add_implementation_approach(self, doc, concept: Dict[str, Any]):
        """Add implementation approach section"""
        doc.add_paragraph("4. Implementation Approach", style='CustomHeading')
        doc.add_paragraph("This section describes the implementation approach and methodology.", style='CustomNormal')
        doc.add_paragraph()

    def _add_metadata(self, doc, concept: Dict[str, Any]):
        """Add metadata section"""
        doc.add_paragraph("Metadata", style='CustomHeading')
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='CustomNormal')
        doc.add_paragraph(f"AI Provider: {concept.get('provider', 'Unknown')}", style='CustomNormal')

    def _remove_dot_blocks(self, text: str) -> str:
        """Remove DOT code blocks from text"""
        self.logger.debug("Removing DOT blocks from text")
        self.logger.debug("Input text length: %d", len(text) if text else 0)
        
        if not text:
            return text
        
        # Remove ```dot ... ``` blocks
        text = re.sub(r'```dot[\s\S]+?```', '', text, flags=re.IGNORECASE)
        # Remove ```graph ... ``` blocks
        text = re.sub(r'```graph[\s\S]+?```', '', text, flags=re.IGNORECASE)
        # Remove digraph blocks
        text = re.sub(r'```digraph[\s\S]+?```', '', text, flags=re.IGNORECASE)
        
        # Remove lines that start with 'digraph' or 'graph'
        lines = text.splitlines()
        cleaned_lines = [line for line in lines if not line.strip().lower().startswith(('digraph', 'graph'))]
        result = '\n'.join(cleaned_lines)
        
        self.logger.debug("Text after DOT removal length: %d", len(result))
        return result 